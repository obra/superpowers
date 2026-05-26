#!/usr/bin/env python3
"""Extract the visual + structural shape of a .docx into JSON for builder generation.

Reads a Word document prepared with mail-merge fields and emits an ordered,
generation-ready JSON tree describing sections, paragraphs, runs (split into
static text vs. dynamic field references), and tables.

Usage:
    python3 extract_docx.py <path-to-docx> [--out file.json]

Dependency:
    pip install python-docx

Output goes to stdout (or --out). All measurements stay in the raw OOXML units
that the `docx` npm library consumes directly:
  - sizeHalfPoints: w:sz value (half-points). docx `size` field uses half-points.
  - DXA (twentieths of a point) for spacing, indent, widths, margins.
  - color: 6-char hex without '#', matching docx `color`.
"""

import argparse
import json
import re
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
except ModuleNotFoundError:
    sys.stderr.write(
        "ERROR: python-docx not installed.\n"
        "Install it with:  pip install python-docx\n"
    )
    sys.exit(2)


# Maps OOXML w:jc values to docx AlignmentType member names.
ALIGNMENT_MAP = {
    "left": "LEFT",
    "center": "CENTER",
    "right": "RIGHT",
    "both": "JUSTIFIED",
    "distribute": "DISTRIBUTE",
}

MERGEFIELD_RE = re.compile(r"MERGEFIELD\s+\"?([^\\\"]+?)\"?\s*(\\.*)?$")
# Plain-text mail-merge tokens typed directly, e.g. «passengerName».
GUILLEMET_RE = re.compile(r"«\s*([^»]+?)\s*»")


def _attr(el, name):
    return el.get(qn(name)) if el is not None else None


def _bool_prop(rpr, tag):
    """A toggle property (w:b, w:i, w:u) is on unless explicitly val=false/0."""
    if rpr is None:
        return False
    el = rpr.find(qn(tag))
    if el is None:
        return False
    val = _attr(el, "w:val")
    return val not in ("false", "0", "off")


def run_props(rpr):
    """Extract run formatting from a w:rPr element into a generation-ready dict."""
    props = {
        "bold": _bool_prop(rpr, "w:b"),
        "italics": _bool_prop(rpr, "w:i"),
        "underline": False,
        "font": None,
        "sizeHalfPoints": None,
        "color": None,
        "allCaps": _bool_prop(rpr, "w:caps"),
    }
    if rpr is None:
        return props

    u = rpr.find(qn("w:u"))
    if u is not None and _attr(u, "w:val") not in ("none", None):
        props["underline"] = True

    fonts = rpr.find(qn("w:rFonts"))
    if fonts is not None:
        props["font"] = _attr(fonts, "w:ascii") or _attr(fonts, "w:hAnsi")

    sz = rpr.find(qn("w:sz"))
    if sz is not None:
        val = _attr(sz, "w:val")
        if val and val.isdigit():
            props["sizeHalfPoints"] = int(val)

    color = rpr.find(qn("w:color"))
    if color is not None:
        val = _attr(color, "w:val")
        if val and val != "auto":
            props["color"] = val

    return props


def paragraph_props(ppr):
    """Extract paragraph formatting from a w:pPr element."""
    props = {
        "style": None,
        "alignment": None,
        "spacing": None,
        "indent": None,
        "numbered": False,
    }
    if ppr is None:
        return props

    style = ppr.find(qn("w:pStyle"))
    if style is not None:
        props["style"] = _attr(style, "w:val")

    jc = ppr.find(qn("w:jc"))
    if jc is not None:
        props["alignment"] = ALIGNMENT_MAP.get(_attr(jc, "w:val"))

    spacing = ppr.find(qn("w:spacing"))
    if spacing is not None:
        sp = {}
        for k in ("w:before", "w:after", "w:line"):
            v = _attr(spacing, k)
            if v is not None and v.lstrip("-").isdigit():
                sp[k.split(":")[1]] = int(v)
        if sp:
            props["spacing"] = sp

    ind = ppr.find(qn("w:ind"))
    if ind is not None:
        indent = {}
        for k in ("w:left", "w:right", "w:firstLine", "w:hanging"):
            v = _attr(ind, k)
            if v is not None and v.lstrip("-").isdigit():
                indent[k.split(":")[1]] = int(v)
        if indent:
            props["indent"] = indent

    if ppr.find(qn("w:numPr")) is not None:
        props["numbered"] = True

    return props


def _first_text(run_el):
    t = run_el.find(qn("w:t"))
    return t.text or "" if t is not None else ""


def parse_paragraph(p_el, fields):
    """Walk a w:p element in document order, producing ordered run segments.

    Each segment is either:
      {"kind": "static", "text": ..., <run props>}
      {"kind": "field",  "field": <name>, <run props>}

    Complex mail-merge fields (fldChar begin/separate/end) and simple fields
    (w:fldSimple) are recognized, as well as plain-text «token» guillemets.
    """
    props = paragraph_props(p_el.find(qn("w:pPr")))
    segments = []

    # Field state machine for complex fields.
    in_field_instr = False
    in_field_result = False
    instr_buffer = ""
    field_props = None

    def emit_static(text, rpr):
        if not text:
            return
        # Split out any inline «token» guillemets typed as plain text.
        last = 0
        for m in GUILLEMET_RE.finditer(text):
            if m.start() > last:
                seg = run_props(rpr)
                seg.update({"kind": "static", "text": text[last:m.start()]})
                segments.append(seg)
            name = m.group(1).strip()
            _register(name, rpr)
            last = m.end()
        if last < len(text):
            seg = run_props(rpr)
            seg.update({"kind": "static", "text": text[last:]})
            segments.append(seg)

    def _register(name, rpr):
        if name not in fields:
            fields.append(name)
        seg = run_props(rpr)
        seg.update({"kind": "field", "field": name})
        segments.append(seg)

    def handle_run(run_el):
        nonlocal in_field_instr, in_field_result, instr_buffer, field_props
        rpr = run_el.find(qn("w:rPr"))
        for child in run_el:
            tag = child.tag
            if tag == qn("w:fldChar"):
                ft = _attr(child, "w:fldCharType")
                if ft == "begin":
                    in_field_instr = True
                    instr_buffer = ""
                    field_props = rpr
                elif ft == "separate":
                    in_field_instr = False
                    in_field_result = True
                elif ft == "end":
                    in_field_instr = False
                    in_field_result = False
                    name = _parse_instr(instr_buffer)
                    if name:
                        _register(name, field_props or rpr)
                    instr_buffer = ""
            elif tag == qn("w:instrText"):
                if in_field_instr:
                    instr_buffer += child.text or ""
            elif tag == qn("w:t"):
                if in_field_result:
                    # Cached field display value — skip, it is regenerated.
                    continue
                if in_field_instr:
                    continue
                emit_static(child.text or "", rpr)
            elif tag == qn("w:tab"):
                if not (in_field_instr or in_field_result):
                    seg = run_props(rpr)
                    seg.update({"kind": "static", "text": "\t"})
                    segments.append(seg)

    def _parse_instr(buf):
        m = MERGEFIELD_RE.search(buf.strip())
        return m.group(1).strip() if m else None

    for child in p_el:
        tag = child.tag
        if tag == qn("w:r"):
            handle_run(child)
        elif tag == qn("w:fldSimple"):
            instr = _attr(child, "w:instr") or ""
            name = _parse_instr(instr)
            inner_rpr = None
            r = child.find(qn("w:r"))
            if r is not None:
                inner_rpr = r.find(qn("w:rPr"))
            if name:
                _register(name, inner_rpr)
        elif tag == qn("w:hyperlink"):
            for r in child.findall(qn("w:r")):
                handle_run(r)

    return {"type": "paragraph", **props, "runs": segments}


def parse_table(tbl_el, fields):
    grid = []
    grid_el = tbl_el.find(qn("w:tblGrid"))
    if grid_el is not None:
        for col in grid_el.findall(qn("w:gridCol")):
            w = _attr(col, "w:w")
            grid.append(int(w) if w and w.isdigit() else 0)

    borders = False
    tblpr = tbl_el.find(qn("w:tblPr"))
    if tblpr is not None and tblpr.find(qn("w:tblBorders")) is not None:
        borders = True

    rows = []
    for tr in tbl_el.findall(qn("w:tr")):
        cells = []
        for tc in tr.findall(qn("w:tc")):
            tcpr = tc.find(qn("w:tcPr"))
            cell = {"gridSpan": 1, "vMerge": None, "shading": None, "width": None,
                    "content": []}
            if tcpr is not None:
                gs = tcpr.find(qn("w:gridSpan"))
                if gs is not None:
                    v = _attr(gs, "w:val")
                    cell["gridSpan"] = int(v) if v and v.isdigit() else 1
                vm = tcpr.find(qn("w:vMerge"))
                if vm is not None:
                    cell["vMerge"] = _attr(vm, "w:val") or "continue"
                shd = tcpr.find(qn("w:shd"))
                if shd is not None:
                    fill = _attr(shd, "w:fill")
                    if fill and fill != "auto":
                        cell["shading"] = fill
                tcw = tcpr.find(qn("w:tcW"))
                if tcw is not None:
                    w = _attr(tcw, "w:w")
                    if w and w.isdigit():
                        cell["width"] = int(w)
            for p in tc.findall(qn("w:p")):
                cell["content"].append(parse_paragraph(p, fields))
            cells.append(cell)
        rows.append({"cells": cells})

    return {"type": "table", "borders": borders, "columnWidths": grid, "rows": rows}


def parse_section(sectpr):
    section = {"page": None, "headerRefs": [], "footerRefs": []}
    if sectpr is None:
        return section

    pgsz = sectpr.find(qn("w:pgSz"))
    pgmar = sectpr.find(qn("w:pgMar"))
    page = {}
    if pgsz is not None:
        for k in ("w:w", "w:h"):
            v = _attr(pgsz, k)
            if v and v.isdigit():
                page["width" if k == "w:w" else "height"] = int(v)
    if pgmar is not None:
        margins = {}
        for k in ("w:top", "w:right", "w:bottom", "w:left", "w:header", "w:footer"):
            v = _attr(pgmar, k)
            if v is not None and v.lstrip("-").isdigit():
                margins[k.split(":")[1]] = int(v)
        if margins:
            page["margins"] = margins
    if page:
        section["page"] = page

    for ref in sectpr.findall(qn("w:headerReference")):
        section["headerRefs"].append(_attr(ref, "w:type"))
    for ref in sectpr.findall(qn("w:footerReference")):
        section["footerRefs"].append(_attr(ref, "w:type"))

    return section


def extract(path):
    doc = Document(path)
    fields = []
    body = []

    body_el = doc.element.body
    for child in body_el:
        if child.tag == qn("w:p"):
            body.append(parse_paragraph(child, fields))
        elif child.tag == qn("w:tbl"):
            body.append(parse_table(child, fields))

    sectpr = body_el.find(qn("w:sectPr"))
    section = parse_section(sectpr)

    # Header/footer text presence (helps decide buildInstitutionalHeader usage).
    has_header = bool(doc.sections and doc.sections[0].header.paragraphs
                      and any(p.text.strip() for p in doc.sections[0].header.paragraphs)) \
        if doc.sections else False
    has_footer = bool(doc.sections and doc.sections[0].footer.paragraphs
                      and any(p.text.strip() for p in doc.sections[0].footer.paragraphs)) \
        if doc.sections else False

    return {
        "source": path,
        "fields": fields,
        "section": section,
        "hasHeaderContent": has_header,
        "hasFooterContent": has_footer,
        "body": body,
    }


def main():
    parser = argparse.ArgumentParser(description="Extract .docx structure to JSON.")
    parser.add_argument("docx", help="Path to the .docx file")
    parser.add_argument("--out", help="Write JSON to this file instead of stdout")
    args = parser.parse_args()

    result = extract(args.docx)
    payload = json.dumps(result, ensure_ascii=False, indent=2)

    if args.out:
        with open(args.out, "w", encoding="utf-8") as fh:
            fh.write(payload)
        sys.stderr.write(f"Wrote {args.out}\n")
    else:
        sys.stdout.write(payload + "\n")


if __name__ == "__main__":
    main()
