#!/usr/bin/env python3
"""
prep_doc.py — Tiền xử lý tài liệu PO trước khi review.

Làm phần CƠ HỌC (rẻ, ổn định) cho quy trình "doc dài":
  1) Lấy text sạch (docx / export Confluence MIME / html / md / txt)
  2) Đo kích thước CHỮ THẬT (chars / words / ~tokens) — quyết định có cần quy trình doc dài không
  3) Tách section theo heading (best-effort) → ghi từng mục ra file để review riêng
  4) Trích "chỉ mục claim" để PASS ĐỐI SOÁT bắt mâu thuẫn xa nhau:
       - link Figma
       - các con số đếm ("13 nhóm", "20/trang", "12/24 buổi"...) kèm ngữ cảnh + section

Việc PHÁN ĐOÁN (mơ hồ, logic, enum/glossary, doc↔figma) vẫn do model làm — script chỉ dọn đường.

Usage:
    python prep_doc.py <path-to-doc> [out_dir]
Mặc định out_dir = ./prep_out
"""

import sys, os, re, json, email, html, argparse
from pathlib import Path
from html.parser import HTMLParser

# ---------- 1) Lấy text sạch ----------

class _Text(HTMLParser):
    def __init__(self):
        super().__init__(); self.out=[]; self.skip=0
    def handle_starttag(self, tag, attrs):
        if tag in ('script','style','head'): self.skip+=1
        if tag in ('tr','br','p','div','li','h1','h2','h3','h4','h5','h6','td','th'): self.out.append('\n')
        if tag in ('td','th'): self.out.append(' | ')
    def handle_endtag(self, tag):
        if tag in ('script','style','head'): self.skip=max(0,self.skip-1)
    def handle_data(self, data):
        if self.skip==0:
            t=data.strip()
            if t: self.out.append(t+' ')

def _strip_html(s: str) -> str:
    p=_Text(); p.feed(s); return ''.join(p.out)

def extract_text(path: Path) -> str:
    raw = path.read_bytes()
    suf = path.suffix.lower()
    head = raw[:200].lstrip()
    # Confluence "Word" export = MIME multipart (đuôi .doc/.mht/.eml nhưng thực ra là email)
    if head[:5] in (b'Date:', b'MIME-', b'From:', b'Conte') or suf in ('.mht','.mhtml','.eml'):
        try:
            msg = email.message_from_bytes(raw)
            for part in msg.walk():
                if part.get_content_type()=='text/html':
                    return _strip_html(part.get_payload(decode=True).decode('utf-8','replace'))
            for part in msg.walk():
                if part.get_content_type()=='text/plain':
                    return part.get_payload(decode=True).decode('utf-8','replace')
        except Exception:
            pass
    if suf in ('.html','.htm') or head[:5].lower()==b'<html':
        return _strip_html(raw.decode('utf-8','replace'))
    if suf in ('.docx','.dotx'):
        try:
            import docx  # python-docx
            d=docx.Document(str(path))
            parts=[]
            for p in d.paragraphs: parts.append(p.text)
            for t in d.tables:
                for r in t.rows: parts.append(' | '.join(c.text for c in r.cells))
            return '\n'.join(parts)
        except Exception:
            import subprocess
            try:
                return subprocess.run(['pandoc',str(path),'-t','plain'],capture_output=True,text=True,check=True).stdout
            except Exception as e:
                raise SystemExit(f"Cần python-docx hoặc pandoc để đọc docx: {e}")
    # md / txt / fallback
    return raw.decode('utf-8','replace')

def normalize(text: str) -> str:
    text = html.unescape(text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    return text.strip()

# ---------- 3) Tách section (best-effort) ----------

def _upper_ratio(s: str) -> float:
    letters=[c for c in s if c.isalpha()]
    return (sum(c.isupper() for c in letters)/len(letters)) if letters else 0.0

def is_heading(line: str) -> bool:
    s=line.strip()
    if not (3 <= len(s) <= 90): return False
    if re.match(r'^US-[A-Za-z0-9][\w.\-]*', s): return True
    if re.match(r'^#{1,6}\s+\S', s): return True
    if re.match(r'^[A-Z]\.\s+\S', s): return True               # A. B. C.
    m=re.match(r'^(\d+(?:\.\d+)*)[\.\)]?\s+(.{1,80})$', s)      # numbered
    if m:
        top = '.' not in m.group(1)                            # top-level "1." "5."
        capsy = _upper_ratio(s) > 0.5
        shortish = len(s.split()) <= 12
        return top or capsy or shortish
    if _upper_ratio(s) > 0.7 and len(s.split()) <= 12:          # ALL CAPS title
        return True
    return False

def split_sections(text: str):
    lines=text.split('\n')
    sections=[]; cur={'title':'(mở đầu)','lines':[]}
    for ln in lines:
        if is_heading(ln):
            if cur['lines']: sections.append(cur)
            cur={'title':ln.strip(),'lines':[]}
        else:
            cur['lines'].append(ln)
    if cur['lines'] or not sections: sections.append(cur)
    for s in sections:
        s['text']='\n'.join(s['lines']).strip()
        s['words']=len(s['text'].split())
        del s['lines']
    return sections

# ---------- 4) Chỉ mục claim ----------

FIGMA_RE = re.compile(r'https?://(?:www\.)?figma\.com/(?:design|file|proto)/[^\s"\'<>]+')
COUNT_RE = re.compile(r'(\d+)\s*(nhóm|cột|tab|bước|trường|trạng thái|mục|buổi|trang|item|ngày|màn|tiêu chí)\b', re.IGNORECASE)
SLASH_RE = re.compile(r'(?<![\d/])(\d{1,3})\s*/\s*(\d{1,3})(?![\d/])')  # "12/24" (loại ngày dd/mm/yyyy)
PAGE_RE  = re.compile(r'\[\s*\d+(?:\s*,\s*\d+)+\s*\]')          # "[20, 50, 100]"

def section_of(idx, sec_offsets):
    title='(mở đầu)'
    for off,t in sec_offsets:
        if off<=idx: title=t
        else: break
    return title

def claim_index(text: str, sections):
    # offsets của heading trong text gốc để gắn claim vào section
    offs=[]; pos=0
    for s in sections:
        i=text.find(s['title'], pos) if s['title']!='(mở đầu)' else 0
        if i<0: i=pos
        offs.append((i, s['title'])); pos=max(pos, i+1)
    def ctx(i,j,pad=45): return re.sub(r'\s+',' ',text[max(0,i-pad):j+pad]).strip()
    figma=sorted(set(FIGMA_RE.findall(text)))
    counts=[]
    for m in COUNT_RE.finditer(text):
        counts.append({'value':m.group(1),'unit':m.group(2).lower(),
                       'section':section_of(m.start(),offs),'context':ctx(*m.span())})
    slashes=[{'value':f"{m.group(1)}/{m.group(2)}",'section':section_of(m.start(),offs),
              'context':ctx(*m.span())} for m in SLASH_RE.finditer(text)]
    pages=[{'value':m.group(0),'section':section_of(m.start(),offs),
            'context':ctx(*m.span())} for m in PAGE_RE.finditer(text)]
    return {'figma_links':figma,'count_claims':counts,'ratio_claims':slashes,'page_size_claims':pages}

# ---------- main ----------

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument('path'); ap.add_argument('out_dir', nargs='?', default='prep_out')
    a=ap.parse_args()
    p=Path(a.path)
    if not p.exists(): raise SystemExit(f"Không thấy file: {p}")
    text=normalize(extract_text(p))
    out=Path(a.out_dir); (out/'sections').mkdir(parents=True, exist_ok=True)

    words=len(text.split()); chars=len(text); toks=chars//4
    secs=split_sections(text)
    idx=claim_index(text, secs)

    # ghi text sạch + từng section
    (out/'clean.txt').write_text(text, encoding='utf-8')
    for n,s in enumerate(secs,1):
        slug=re.sub(r'[^\w]+','_', s['title'])[:40].strip('_') or 'mucs'
        (out/'sections'/f"{n:02d}_{slug}.txt").write_text(
            f"# {s['title']}\n\n{s['text']}", encoding='utf-8')
    (out/'index.json').write_text(json.dumps(idx, ensure_ascii=False, indent=2), encoding='utf-8')

    # gợi ý chiến lược
    us_count = len(re.findall(r'(?m)^\s*US-[A-Za-z0-9]', text))
    LONG = words>6000 or us_count>3
    print(f"== prep_doc: {p.name} ==")
    print(f"Chữ sạch: {chars:,} ký tự · {words:,} từ · ~{toks:,} tokens")
    print(f"Số section phát hiện: {len(secs)}  →  {out}/sections/")
    print(f"Chiến lược đề xuất: {'DOC DÀI → map→deep→đối soát + ngân sách output' if LONG else 'ngắn → review 1 lượt bình thường'}")
    print("\nOutline (section · số từ):")
    for n,s in enumerate(secs,1):
        print(f"  {n:02d}. {s['title'][:70]}  ({s['words']} từ)")
    print(f"\nLink Figma: {len(idx['figma_links'])}")
    for l in idx['figma_links']: print(f"  - {l}")
    if idx['count_claims']:
        print(f"\nCon số đếm (đối soát chéo — số khớp nhau giữa các mục không?):")
        for c in idx['count_claims']:
            print(f"  - {c['value']} {c['unit']}  [{c['section'][:30]}]  «…{c['context']}…»")
    if idx['ratio_claims']:
        print("\nTỉ lệ X/Y:", ', '.join(c['value'] for c in idx['ratio_claims']))
    if idx['page_size_claims']:
        print("Page-size:", ', '.join(c['value'] for c in idx['page_size_claims']))
    print(f"\n→ Đã ghi: {out}/clean.txt, {out}/index.json, {out}/sections/*.txt")

if __name__=='__main__':
    main()
